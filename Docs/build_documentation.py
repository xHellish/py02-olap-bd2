#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Docs/build_documentation.py

Genera la documentacion tecnica editable del Proyecto 2 (Analisis y Procesamiento
OLAP) como un archivo .docx en el root del repositorio.

Tambien renderiza 4 diagramas (arquitectura, esquema estrella, modelo de grafo y
DAG de Airflow) en PNG con matplotlib y los incrusta en el documento.

Uso:
    python Docs/build_documentation.py

Salida:
    Documentacion_Tecnica_PY02_OLAP.docx   (en el root)
    Docs/diagrams/arquitectura.png
    Docs/diagrams/star_schema.png
    Docs/diagrams/graph_model.png
    Docs/diagrams/airflow_dag.png
"""
import os

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
#  Rutas
# ─────────────────────────────────────────────────────────────────────────────
HERE     = os.path.dirname(os.path.abspath(__file__))
ROOT     = os.path.dirname(HERE)
DIAGRAMS = os.path.join(HERE, "diagrams")
OUT_DOCX = os.path.join(ROOT, "Documentacion_Tecnica_PY02_OLAP.docx")
os.makedirs(DIAGRAMS, exist_ok=True)

# Paleta
NAVY   = "#1F3864"
BLUE   = "#2E5496"
LIGHT  = "#D9E1F2"
ACCENT = "#C55A11"
GREEN  = "#548235"
GREY   = "#808080"

# ═════════════════════════════════════════════════════════════════════════════
#  1. DIAGRAMAS (matplotlib → PNG)
# ═════════════════════════════════════════════════════════════════════════════

def _box(ax, x, y, w, h, text, fc, ec=NAVY, tc="white", fs=9, weight="bold"):
    ax.add_patch(FancyBboxPatch((x, y), w, h,
                 boxstyle="round,pad=0.02,rounding_size=0.06",
                 linewidth=1.4, edgecolor=ec, facecolor=fc, zorder=2))
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
            fontsize=fs, color=tc, weight=weight, zorder=3, wrap=True)


def _arrow(ax, x1, y1, x2, y2, color=GREY, style="-|>", lw=1.6, ls="-"):
    ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2),
                 arrowstyle=style, mutation_scale=14, linewidth=lw,
                 color=color, linestyle=ls, zorder=1,
                 connectionstyle="arc3,rad=0.0"))


def diagram_architecture():
    fig, ax = plt.subplots(figsize=(11, 6.2))
    ax.set_xlim(0, 22); ax.set_ylim(0, 12); ax.axis("off")

    # Capa fuente (OLTP)
    _box(ax, 0.5, 9.0, 4.2, 1.4, "PostgreSQL OLTP\nrestaurantes_oltp", BLUE)
    _box(ax, 0.5, 7.2, 4.2, 1.2, "MongoDB", "#6B7DA3")
    _box(ax, 0.5, 5.4, 4.2, 1.2, "Elasticsearch", "#6B7DA3")
    ax.text(2.6, 10.7, "FUENTES (CORE)", ha="center", fontsize=9, weight="bold", color=NAVY)

    # Spark
    _box(ax, 6.0, 7.0, 4.0, 2.4, "Apache Spark\n6 jobs ETL\n(DataFrame API\n+ SparkSQL)", ACCENT)
    ax.text(8.0, 9.7, "PROCESAMIENTO", ha="center", fontsize=9, weight="bold", color=NAVY)

    # Hive DW
    _box(ax, 11.3, 8.4, 4.2, 1.5, "Hive Data Warehouse\nrestaurantes_dw\n(ORC + Snappy)", NAVY)
    _box(ax, 11.3, 6.5, 4.2, 1.4, "8 dimensiones\n2 hechos + 6 cubos OLAP", "#3A5A98")
    ax.text(13.4, 10.2, "DATA WAREHOUSE", ha="center", fontsize=9, weight="bold", color=NAVY)

    # Serving + Superset
    _box(ax, 16.8, 8.4, 4.6, 1.5, "PostgreSQL Serving\nrestaurantes_marts\n(10+ marts)", GREEN)
    _box(ax, 16.8, 6.5, 4.6, 1.4, "Apache Superset\n3 dashboards / 9 charts", "#6FA64C")
    ax.text(19.1, 10.2, "VISUALIZACION", ha="center", fontsize=9, weight="bold", color=NAVY)

    # Neo4J + routing
    _box(ax, 11.3, 3.4, 4.2, 1.5, "Neo4J + GDS\n6 nodos / 7 relaciones", "#7030A0")
    _box(ax, 16.8, 3.4, 4.6, 1.5, "Routing\nvecino mas cercano\n(route_results.json)", "#9A5BC0")
    ax.text(13.4, 5.2, "GRAFOS Y RUTAS", ha="center", fontsize=9, weight="bold", color=NAVY)

    # Airflow (orquestador, abajo)
    _box(ax, 1.5, 1.1, 19.0, 1.3,
         "Apache Airflow   ·   DAG etl_olap_dag (10 tareas, @daily)   ·   orquesta el pipeline ETL completo",
         "#385723", fs=9.5)

    # Flechas principales
    _arrow(ax, 4.7, 9.6, 6.0, 8.6)            # OLTP -> Spark
    _arrow(ax, 4.7, 7.7, 6.0, 8.0)            # Mongo -> Spark
    _arrow(ax, 10.0, 8.4, 11.3, 9.0)          # Spark -> Hive DW
    _arrow(ax, 10.0, 7.6, 11.3, 7.2)          # Spark -> dims/facts
    _arrow(ax, 15.5, 9.0, 16.8, 9.0)          # Hive -> serving (export_marts)
    _arrow(ax, 19.1, 8.4, 19.1, 7.9)          # serving -> superset
    _arrow(ax, 4.7, 6.0, 11.3, 4.1, color="#7030A0")  # OLTP -> Neo4J (CSV)
    _arrow(ax, 15.5, 4.1, 16.8, 4.1, color="#9A5BC0") # Neo4J -> routing
    # Airflow controla (una sola flecha clara hacia la capa de procesamiento)
    _arrow(ax, 8.0, 2.4, 8.0, 6.9, color="#385723", lw=1.2, ls=(0, (4, 3)))

    ax.text(11.0, 0.4, "Despliegue: Docker Compose (perfiles core, dw, orchestration, graph, viz) sobre red olap-net",
            ha="center", fontsize=8, style="italic", color=GREY)

    fig.tight_layout()
    out = os.path.join(DIAGRAMS, "arquitectura.png")
    fig.savefig(out, dpi=170, bbox_inches="tight"); plt.close(fig)
    print(f"  diagrama arquitectura -> {out}")


def diagram_star_schema():
    fig, ax = plt.subplots(figsize=(11, 7.4))
    ax.set_xlim(0, 22); ax.set_ylim(0, 15); ax.axis("off")

    # Hechos al centro
    _box(ax, 8.4, 8.2, 5.2, 1.8, "FACT_ORDERS\n(1 fila / pedido)\npart. year, month", ACCENT, fs=9)
    _box(ax, 8.4, 5.0, 5.2, 1.8, "FACT_ORDER_ITEMS\n(1 fila / item)\npart. year, month", ACCENT, fs=9)

    dims = [
        ("DIM_DATE\n(365 filas)",       1.0, 12.6),
        ("DIM_TIME\n(24 filas)",        8.4, 12.6),
        ("DIM_CATEGORY",               15.8, 12.6),
        ("DIM_PRODUCT\n(SCD Tipo 1)",  18.6,  8.2),
        ("DIM_CUSTOMER",               18.6,  5.0),
        ("DIM_RESTAURANT",             15.8,  1.4),
        ("DIM_LOCATION\n(7 zonas)",     8.4,  1.4),
        ("DIM_COURIER",                 1.0,  1.4),
    ]
    coords = {}
    for name, x, y in dims:
        _box(ax, x, y, 3.6, 1.5, name, NAVY, fs=8.5)
        coords[name.split("\n")[0]] = (x + 1.8, y + 0.75)

    fo = (11.0, 9.1)
    fi = (11.0, 5.9)
    # FK fact_orders
    for d in ["DIM_DATE", "DIM_TIME", "DIM_CUSTOMER", "DIM_RESTAURANT",
              "DIM_LOCATION", "DIM_COURIER"]:
        _arrow(ax, coords[d][0], coords[d][1], fo[0], fo[1], color=BLUE, style="-", lw=1.1)
    # FK fact_order_items
    for d in ["DIM_DATE", "DIM_TIME", "DIM_PRODUCT", "DIM_CATEGORY",
              "DIM_CUSTOMER", "DIM_RESTAURANT"]:
        _arrow(ax, coords[d][0], coords[d][1], fi[0], fi[1], color=GREEN, style="-", lw=1.1)

    ax.text(11.0, 14.4, "ESQUEMA ESTRELLA — restaurantes_dw",
            ha="center", fontsize=11, weight="bold", color=NAVY)
    ax.text(0.6, 0.4, "Lineas azules: FK de fact_orders   ·   Lineas verdes: FK de fact_order_items",
            ha="left", fontsize=8, style="italic", color=GREY)

    fig.tight_layout()
    out = os.path.join(DIAGRAMS, "star_schema.png")
    fig.savefig(out, dpi=170, bbox_inches="tight"); plt.close(fig)
    print(f"  diagrama estrella -> {out}")


def diagram_graph_model():
    import math
    fig, ax = plt.subplots(figsize=(11, 7.0))
    ax.set_xlim(0, 22); ax.set_ylim(-1.8, 14); ax.axis("off")

    nodes = {
        "User":       (2.0, 9.5, "#2E5496"),
        "Order":      (9.0, 9.5, ACCENT),
        "Restaurant": (16.0, 9.5, "#2E5496"),
        "Product":    (16.0, 4.0, GREEN),
        "Courier":    (2.0, 4.0, "#7030A0"),
        "Location":   (9.0, 2.0, "#385723"),
    }
    R = 1.5
    for name, (x, y, c) in nodes.items():
        ax.add_patch(plt.Circle((x, y), R, color=c, zorder=2))
        ax.text(x, y, name, ha="center", va="center", color="white",
                fontsize=9.5, weight="bold", zorder=3)

    def edge(a, b, label, rad=0.0, color=GREY):
        (x1, y1, _), (x2, y2, _) = nodes[a], nodes[b]
        ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>",
                     mutation_scale=14, lw=1.5, color=color, zorder=1,
                     connectionstyle=f"arc3,rad={rad}",
                     shrinkA=38, shrinkB=38))
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my + rad * 6 + 0.35, label, ha="center", fontsize=8,
                weight="bold", color=color,
                bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="none", alpha=0.85))

    def self_loop(cx, cy, color, label, above=True):
        sign = 1 if above else -1
        a1, a2 = (62, 118) if above else (-62, -118)
        sx = cx + R * math.cos(math.radians(a1)); sy = cy + R * math.sin(math.radians(a1))
        ex = cx + R * math.cos(math.radians(a2)); ey = cy + R * math.sin(math.radians(a2))
        ax.add_patch(FancyArrowPatch((sx, sy), (ex, ey), arrowstyle="-|>",
                     mutation_scale=13, lw=1.5, color=color, zorder=1,
                     connectionstyle=f"arc3,rad={1.4 * sign}"))
        ax.text(cx, cy + sign * (R + 1.7), label, ha="center", fontsize=8.5,
                weight="bold", color=color)

    edge("User", "Order", "PLACED")
    edge("Order", "Restaurant", "FROM_RESTAURANT")
    edge("Order", "Product", "CONTAINS\n{qty, price}", rad=0.12)
    edge("Order", "Location", "DELIVERED_TO\n{dist, eta}", rad=-0.12)
    edge("Courier", "Order", "ASSIGNED", rad=0.12)

    # RECOMMENDS: loop sobre el nodo User
    ux, uy, _ = nodes["User"]
    self_loop(ux, uy, "#9A5BC0", "RECOMMENDS (User → User)", above=True)
    # ROUTE: loop bajo el nodo Location
    lx, ly, _ = nodes["Location"]
    self_loop(lx, ly, "#385723", "ROUTE (Location ↔ Location)\n{distance_km, time_min}", above=False)

    ax.text(11.0, 13.4, "MODELO DE GRAFO — Neo4J", ha="center",
            fontsize=11, weight="bold", color=NAVY)
    fig.tight_layout()
    out = os.path.join(DIAGRAMS, "graph_model.png")
    fig.savefig(out, dpi=170, bbox_inches="tight"); plt.close(fig)
    print(f"  diagrama grafo -> {out}")


def diagram_airflow_dag():
    fig, ax = plt.subplots(figsize=(11, 5.4))
    ax.set_xlim(0, 24); ax.set_ylim(0, 11); ax.axis("off")

    W, H = 3.6, 1.1
    def t(x, y, label, c=BLUE):
        _box(ax, x, y, W, H, label, c, fs=7.8)
        return (x, y)

    t(0.3, 5.0, "extract_postgres")
    t(4.4, 5.0, "build_hive_dw")
    t(8.7, 8.4, "run_trends")
    t(8.7, 6.6, "run_peak_hours")
    t(8.7, 4.8, "run_monthly_growth")
    t(8.7, 2.6, "check_catalog_changed", ACCENT)
    t(13.4, 3.6, "reindex_elasticsearch", "#7030A0")
    t(13.4, 1.6, "skip_reindex", GREY)
    t(13.4, 6.6, "export_marts", GREEN)
    t(18.6, 4.6, "data_quality_checks", "#385723")

    def link(x1, y1, x2, y2, color=GREY):
        _arrow(ax, x1 + W, y1 + H / 2, x2, y2 + H / 2, color=color, lw=1.3)

    link(0.3, 5.0, 4.4, 5.0)
    for yy in (8.4, 6.6, 4.8, 2.6):
        link(4.4, 5.0, 8.7, yy)
    for yy in (8.4, 6.6, 4.8):
        link(8.7, yy, 13.4, 6.6, GREEN)
    link(8.7, 2.6, 13.4, 3.6, ACCENT)
    link(8.7, 2.6, 13.4, 1.6, ACCENT)
    link(13.4, 6.6, 18.6, 4.6, "#385723")
    link(13.4, 3.6, 18.6, 4.6, "#385723")
    link(13.4, 1.6, 18.6, 4.6, "#385723")

    ax.text(12.0, 10.2, "DAG etl_olap_dag — 10 tareas", ha="center",
            fontsize=11, weight="bold", color=NAVY)
    ax.text(12.0, 0.2, "Branch naranja: reindexa ES solo si cambia el hash del catalogo. "
            "data_quality_checks usa trigger_rule='all_done'.",
            ha="center", fontsize=8, style="italic", color=GREY)
    fig.tight_layout()
    out = os.path.join(DIAGRAMS, "airflow_dag.png")
    fig.savefig(out, dpi=170, bbox_inches="tight"); plt.close(fig)
    print(f"  diagrama DAG -> {out}")


# ═════════════════════════════════════════════════════════════════════════════
#  2. HELPERS DEL DOCUMENTO
# ═════════════════════════════════════════════════════════════════════════════

def shade(paragraph, fill="F2F3F5"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    r = run._r
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); r.append(f1)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-3" \\h \\z \\u'; r.append(it)
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate"); r.append(f2)
    t = OxmlElement("w:t")
    t.text = "[Tabla de contenido: clic derecho aqui -> Actualizar campos -> Actualizar toda la tabla]"
    r.append(t)
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end"); r.append(f3)


def code_block(doc, text, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    shade(p)
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        if i < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)
    return p


def para(doc, text, italic=False, size=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)


def numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_table(doc, headers, rows, col_widths=None, header_fill=NAVY,
              font_size=9, first_col_bold=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hcells = table.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = ""
        p = hcells[i].paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(font_size)
        set_cell_bg(hcells[i], header_fill.lstrip("#"))
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(font_size)
            if first_col_bold and i == 0:
                run.bold = True
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_image(doc, path, width_in=6.4, caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def screenshot_placeholder(doc, titulo, instruccion, ruta_sugerida):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    shade(p, "FFF2CC")
    r = p.add_run(f"[ CAPTURA DE PANTALLA — {titulo} ]")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x80, 0x60, 0x00)
    r.add_break(WD_BREAK.LINE)
    r2 = p.add_run(f"Insertar aqui: {instruccion}")
    r2.font.size = Pt(8.5); r2.italic = True
    r2.add_break(WD_BREAK.LINE)
    r3 = p.add_run(f"Archivo sugerido: {ruta_sugerida}")
    r3.font.size = Pt(8); r3.font.name = "Consolas"; r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def h1(doc, text):
    doc.add_heading(text, level=1)

def h2(doc, text):
    doc.add_heading(text, level=2)

def h3(doc, text):
    doc.add_heading(text, level=3)


# ═════════════════════════════════════════════════════════════════════════════
#  3. CONSTRUCCION DEL DOCUMENTO
# ═════════════════════════════════════════════════════════════════════════════

def build():
    doc = Document()

    # Estilos base
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for lvl, color, size in [(1, "#000000", 14), (2, "#000000", 13), (3, "#000000", 12)]:
        st = doc.styles[f"Heading {lvl}"]
        st.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        st.font.size = Pt(size)
        st.font.name = "Times New Roman"

    # ── PORTADA ───────────────────────────────────────────────────────────────
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Instituto Tecnologico de Costa Rica")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(NAVY.lstrip("#"))
    for txt, sz in [("Unidad de Ingenieria en Computacion — Sede Central Cartago", 11),
                    ("Base de Datos 2", 11),
                    ("Profesor: Kenneth Obando Rodriguez", 11)]:
        q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = q.add_run(txt); rr.font.size = Pt(sz)

    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Proyecto 2\nAnalisis y Procesamiento OLAP")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor.from_string(NAVY.lstrip("#"))

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Plataforma de analitica de pedidos para una red de restaurantes")
    r.italic = True; r.font.size = Pt(12)

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documentacion Tecnica")
    r.bold = True; r.font.size = Pt(14)
    for txt in ["Estudiante: Ben Farz  (benfarz.bf25@gmail.com)",
                "Fecha de entrega: junio 2026",
                "Repositorio: Proyecto 2 OLAP"]:
        q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.add_run(txt).font.size = Pt(11)

    doc.add_page_break()

    # ── TABLA DE CONTENIDO ──────────────────────────────────────────────────────
    h1(doc, "Tabla de contenido")
    add_toc(doc)
    doc.add_page_break()

    # ── 1. INTRODUCCION ─────────────────────────────────────────────────────────
    h1(doc, "1. Introduccion")

    h2(doc, "1.1 Objetivo")
    para(doc, "Este proyecto implementa una plataforma completa de analisis y procesamiento "
              "OLAP para los datos historicos de una red de restaurantes con servicio de "
              "entrega a domicilio. La solucion cubre el ciclo de vida analitico de extremo "
              "a extremo: ingesta desde la base de datos operacional (OLTP), construccion de "
              "un Data Warehouse dimensional, procesamiento masivo con Apache Spark, "
              "orquestacion del pipeline con Apache Airflow, analisis de relaciones y rutas "
              "con Neo4J, y visualizacion mediante dashboards en Apache Superset. Todos los "
              "componentes son open-source y se despliegan con Docker Compose.")

    h2(doc, "1.2 Alcance y componentes")
    para(doc, "El enunciado define seis componentes; la siguiente tabla resume como se "
              "satisface cada uno y su peso en la evaluacion.")
    add_table(doc,
        ["Componente del enunciado", "Peso", "Implementacion en este proyecto"],
        [
            ["1. Arquitectura OLAP y Data Warehouse", "20%",
             "Esquema estrella en Apache Hive: 8 dimensiones, 2 tablas de hechos y 6 cubos OLAP."],
            ["2. Procesamiento con Apache Spark", "20%",
             "6 jobs PySpark (DataFrame API + SparkSQL) con 3 analisis: tendencias, horarios pico y crecimiento mensual."],
            ["3. Visualizacion de datos", "15%",
             "Apache Superset con 3 dashboards y 9 graficos sobre la capa de marts en PostgreSQL."],
            ["4. Orquestacion con Apache Airflow", "15%",
             "DAG etl_olap_dag con 10 tareas, branching condicional y validaciones de calidad."],
            ["5. Neo4J para grafos y rutas", "15%",
             "Grafo de 6 tipos de nodo y 7 relaciones; consultas Cypher y algoritmos GDS (PageRank, Dijkstra)."],
            ["6. Asignacion de rutas de entrega", "10%",
             "Modulo de enrutamiento con heuristica de vecino mas cercano apoyado en distancias del grafo."],
            ["Documentacion", "5%", "El presente documento."],
        ],
        col_widths=[2.5, 0.6, 3.6], font_size=9, first_col_bold=True)

    h2(doc, "1.3 Contexto: relacion con el Proyecto 1")
    para(doc, "El catalogo del dominio (categorias, productos, restaurantes, menus) se hereda "
              "conceptualmente del Proyecto 1. Sin embargo, el Proyecto 1 no contiene el "
              "historial transaccional necesario para el analisis OLAP (pedidos, items, "
              "entregas, geolocalizacion ni ventana temporal). Por ello, este proyecto incluye "
              "un generador de datos sinteticos que produce una historia realista de 12 meses, "
              "manteniendo el codigo del Proyecto 2 completamente fuera de la carpeta del "
              "Proyecto 1 (que se conserva como referencia de solo lectura).")

    # ── 2. ARQUITECTURA ─────────────────────────────────────────────────────────
    h1(doc, "2. Arquitectura general de la solucion")
    add_image(doc, os.path.join(DIAGRAMS, "arquitectura.png"), 6.6,
              "Figura 1. Flujo de datos de extremo a extremo y orquestacion.")

    h2(doc, "2.1 Flujo de datos")
    para(doc, "El flujo principal es: (1) la base OLTP en PostgreSQL almacena las operaciones "
              "del negocio; (2) Spark extrae esos datos a una zona de staging en Parquet; "
              "(3) Spark construye el esquema estrella en Hive (ORC); (4) tres jobs de analisis "
              "producen data marts; (5) los marts y cubos se exportan a una base PostgreSQL de "
              "serving; (6) Superset consume esa base para los dashboards. En paralelo, los "
              "datos del OLTP se exportan a CSV y se cargan en Neo4J para el analisis de grafos "
              "y el enrutamiento. Apache Airflow orquesta y programa todo el pipeline.")

    h2(doc, "2.2 Componentes y tecnologias")
    add_table(doc,
        ["Servicio", "Imagen / version", "Rol", "Perfil"],
        [
            ["postgres-oltp", "postgres:16-alpine", "Base operacional fuente", "core"],
            ["mongo", "mongo:8", "Almacen documental (heredado P1)", "core"],
            ["elasticsearch", "elasticsearch:8.15.5", "Indice de busqueda de catalogo", "core"],
            ["hive-metastore", "apache/hive:4.0.0", "Metastore (Postgres backend)", "dw"],
            ["hiveserver2", "apache/hive:4.0.0", "Acceso JDBC/Beeline al DW", "dw"],
            ["spark-master / worker", "apache/spark:3.5.1", "Motor de procesamiento ETL", "dw"],
            ["postgres-airflow", "postgres:16-alpine", "Metastore de Airflow", "orchestration"],
            ["airflow-webserver / scheduler", "apache/airflow:2.10.4", "Orquestacion del pipeline", "orchestration"],
            ["neo4j", "neo4j:5.24-community", "Grafo + GDS (rutas/recomendaciones)", "graph"],
            ["postgres-serving", "postgres:16-alpine", "Capa de marts para Superset", "viz"],
            ["superset", "apache/superset:4.1.1", "Dashboards y exploracion", "viz"],
        ],
        col_widths=[1.9, 1.7, 2.2, 0.9], font_size=8.5, first_col_bold=True)

    h2(doc, "2.3 Despliegue con Docker Compose y perfiles")
    para(doc, "Los servicios se agrupan en cinco perfiles (core, dw, orchestration, graph, viz) "
              "para poder levantar solo lo necesario en cada fase y reducir el consumo de "
              "recursos. Todos comparten la red bridge olap-net y usan volumenes nombrados "
              "para persistencia.")
    add_table(doc,
        ["Recurso / UI", "URL o puerto", "Credenciales"],
        [
            ["PostgreSQL OLTP", "localhost:5432  (db restaurantes_oltp)", "olap / olap123"],
            ["PostgreSQL Serving", "localhost:5433  (db restaurantes_marts)", "olap / olap123"],
            ["Spark Master UI", "http://localhost:8080", "-"],
            ["HiveServer2 Web UI", "http://localhost:10002", "usuario hive (sin password)"],
            ["Airflow Web UI", "http://localhost:8081", "admin / admin"],
            ["Neo4J Browser", "http://localhost:7474  (bolt 7687)", "neo4j / neo4j123"],
            ["Apache Superset", "http://localhost:8088", "admin / admin"],
            ["Elasticsearch", "http://localhost:9200", "-"],
        ],
        col_widths=[2.0, 3.0, 1.7], font_size=9, first_col_bold=True)

    # ── 3. MODELO FUENTE Y GENERACION ────────────────────────────────────────────
    h1(doc, "3. Modelo de datos fuente (OLTP) y generacion de datos")

    h2(doc, "3.1 Esquema OLTP")
    para(doc, "La base operacional en PostgreSQL contiene el catalogo del dominio mas las "
              "entidades transaccionales propias del Proyecto 2 (pedidos, items, entregas, "
              "repartidores y geolocalizacion). Tablas principales:")
    add_table(doc,
        ["Tabla", "Descripcion", "Cardinalidad aprox."],
        [
            ["categories", "Categorias de producto", "8"],
            ["restaurants", "Restaurantes con zona y coordenadas", "5"],
            ["products", "Productos con precio y categoria", "32"],
            ["menus / menu_products", "Menus y su composicion", "5 menus"],
            ["users", "Clientes + admin con zona y geo", "81 (80 clientes)"],
            ["couriers", "Repartidores con vehiculo y zona base", "12"],
            ["orders", "Pedidos con estado, total y pago", "~20,000"],
            ["order_items", "Lineas de item por pedido", "~40,000"],
            ["deliveries", "Entrega: zona destino, distancia, ETA", "~15,000"],
        ],
        col_widths=[1.7, 3.4, 1.6], font_size=9, first_col_bold=True)

    h2(doc, "3.2 Generacion de datos sinteticos")
    para(doc, "El script data-generator/generate.py crea una historia coherente y reproducible "
              "(semilla fija = 42). Caracteristicas que dan realismo analitico:")
    bullets(doc, [
        ("Ventana temporal: ", "12 meses (julio 2025 a junio 2026), 20,000 pedidos."),
        ("Tendencia de crecimiento: ", "los pedidos por mes crecen linealmente de un factor 0.6 al 1.4, simulando expansion del negocio."),
        ("Horarios pico: ", "distribucion de probabilidad por hora con picos en almuerzo (12-14h) y cena (18-21h)."),
        ("Estacionalidad semanal: ", "mayor demanda viernes, sabado y domingo."),
        ("Co-compra: ", "grupos de afinidad de productos para que ciertos articulos aparezcan juntos (insumo del analisis de grafos)."),
        ("Estados de pedido: ", "completed 68%, cancelled 15%, delivering 7%, preparing 5%, confirmed 3%, pending 2%."),
        ("Geolocalizacion: ", "7 zonas reales (Ciudad de Guatemala + Antigua) con coordenadas y jitter por cliente."),
    ])
    para(doc, "Comando de generacion (con el perfil core levantado):")
    code_block(doc, "make seed\n# equivalente a:\n#   pip install -r data-generator/requirements.txt\n#   python data-generator/generate.py")

    # ── 4. DW Y CUBOS ────────────────────────────────────────────────────────────
    h1(doc, "4. Data Warehouse y cubos OLAP")
    para(doc, "Corresponde al componente 1 del enunciado (20%). El DW se modela como un "
              "esquema estrella en Apache Hive, con almacenamiento columnar ORC comprimido con "
              "Snappy y tablas de hechos particionadas por anio y mes.")
    add_image(doc, os.path.join(DIAGRAMS, "star_schema.png"), 6.6,
              "Figura 2. Esquema estrella: 2 hechos rodeados de 8 dimensiones.")

    h2(doc, "4.1 Dimensiones")
    add_table(doc,
        ["Dimension", "Granularidad / contenido", "Notas"],
        [
            ["dim_date", "1 fila por dia (365)", "Generada por Spark (sequence de fechas), no se extrae del OLTP."],
            ["dim_time", "1 fila por hora (24)", "Incluye franja: madrugada/manana/almuerzo/tarde/cena/noche."],
            ["dim_category", "Categoria de producto", "Surrogate key + id natural."],
            ["dim_product", "Producto", "SCD Tipo 1 (precio vigente); categoria desnormalizada."],
            ["dim_customer", "Cliente", "Zona y coordenadas para analisis geografico."],
            ["dim_restaurant", "Restaurante", "Zona, coordenadas y rating."],
            ["dim_location", "Zona de entrega (7)", "Centroide lat/lon por zona."],
            ["dim_courier", "Repartidor", "Vehiculo y zona base."],
        ],
        col_widths=[1.5, 2.6, 2.6], font_size=9, first_col_bold=True)

    h2(doc, "4.2 Tablas de hechos")
    bullets(doc, [
        ("fact_orders: ", "una fila por pedido. Medidas: total, item_count, distance_km, eta_min, "
         "delivery_time_min (calculado como delivered_ts - order_ts), is_delivered. "
         "FKs a las 8 dimensiones (location/courier = -1 cuando no hay entrega)."),
        ("fact_order_items: ", "una fila por linea de item. Medidas: quantity, unit_price, line_total. "
         "Permite analisis a nivel de producto sin pasar por fact_orders."),
    ])
    para(doc, "Ambas se particionan por (year, month), lo que acelera las consultas de "
              "tendencia temporal mediante partition pruning.")

    h2(doc, "4.3 Los 6 cubos OLAP")
    para(doc, "Los cubos se implementan como vistas Hive (archivo warehouse/03_olap_cubes.hql). "
              "Cumplen el requisito de analisis agregado por tiempo, ubicacion, tipo de producto "
              "y frecuencia de uso.")
    add_table(doc,
        ["Cubo (vista)", "Dimensiones de analisis", "Metricas principales"],
        [
            ["cube_revenue_month_category", "tiempo x categoria", "revenue, units_sold, order_count, avg_unit_price"],
            ["cube_orders_by_zone", "ubicacion x tiempo", "order_count, delivered, avg_distance, total_revenue"],
            ["cube_order_status", "tiempo x estado x pago", "order_count, revenue, avg_ticket"],
            ["cube_peak_hours", "hora x dia de semana", "order_count, revenue, avg_items"],
            ["cube_product_ranking", "producto x categoria", "units_sold, total_revenue, revenue_per_unit"],
            ["cube_courier_performance", "repartidor", "completados, avg_distance, avg_delay, revenue"],
        ],
        col_widths=[2.4, 1.9, 2.4], font_size=8.5, first_col_bold=True)
    para(doc, "Ejemplo: definicion del cubo de ingresos por mes y categoria.")
    code_block(doc,
        "CREATE VIEW cube_revenue_month_category AS\n"
        "SELECT d.year, d.month, d.month_name,\n"
        "       CONCAT(d.year,'-',LPAD(d.month,2,'0')) AS year_month,\n"
        "       c.category_name,\n"
        "       COUNT(DISTINCT fi.order_id) AS order_count,\n"
        "       SUM(fi.quantity)            AS units_sold,\n"
        "       ROUND(SUM(fi.line_total),2) AS revenue\n"
        "FROM fact_order_items fi\n"
        "JOIN dim_date d     ON fi.date_key = d.date_key\n"
        "JOIN dim_category c ON fi.category_key = c.category_key\n"
        "GROUP BY d.year, d.month, d.month_name, c.category_name;")

    screenshot_placeholder(doc, "DW en Hive — evidencia",
        "captura de beeline o HiveServer2 mostrando 'SHOW TABLES;' en restaurantes_dw (8 dim + 2 hechos + 6 cubos) y/o el resultado de consultar un cubo.",
        "Docs/screenshots/hive_tables.png")

    h2(doc, "4.4 Decision de almacenamiento")
    para(doc, "Se eligio ORC + Snappy por su compresion eficiente, lectura columnar y soporte "
              "nativo de predicate pushdown en Hive y Spark. El metastore de Hive usa PostgreSQL "
              "en lugar de Derby por ser mas robusto y permitir conexiones concurrentes.")

    # ── 5. SPARK ─────────────────────────────────────────────────────────────────
    h1(doc, "5. Procesamiento con Apache Spark")
    para(doc, "Corresponde al componente 2 del enunciado (20%). El pipeline consta de 6 jobs "
              "PySpark que combinan deliberadamente la DataFrame API y SparkSQL, segun lo "
              "solicitado. Todos son idempotentes.")

    h2(doc, "5.1 Pipeline de jobs")
    add_table(doc,
        ["Job", "Responsabilidad", "Tecnica destacada"],
        [
            ["01_extract.py", "Extrae 10 tablas del OLTP a Parquet (staging)", "JDBC read, desacople de extraccion"],
            ["02_build_dw.py", "Construye 8 dimensiones + 2 hechos en Hive", "DataFrame API + SparkSQL, particionado"],
            ["03_trends.py", "Analisis 1: tendencias de consumo", "Window: running total + RANK por mes"],
            ["04_peak_hours.py", "Analisis 2: horarios pico", "Agregacion hora x dia + ranking global"],
            ["05_monthly_growth.py", "Analisis 3: crecimiento mensual", "lag() para variacion % mes a mes"],
            ["06_export_marts.py", "Exporta marts y cubos a postgres-serving", "JDBC write, resumen ejecutivo"],
        ],
        col_widths=[1.7, 2.8, 2.2], font_size=8.5, first_col_bold=True)

    h2(doc, "5.2 Los tres analisis requeridos")
    bullets(doc, [
        ("Tendencias de consumo (03): ", "ingresos y unidades por categoria a lo largo del tiempo, "
         "con ingreso acumulado (running total) y ranking de categoria dentro de cada mes."),
        ("Horarios pico (04): ", "conteo de pedidos por hora x dia de semana para identificar las "
         "franjas de mayor demanda; produce un heatmap operacional y un ranking global."),
        ("Crecimiento mensual (05): ", "variacion porcentual mes a mes de pedidos e ingresos usando "
         "la funcion lag(), mas ingreso acumulado; solo considera pedidos completados."),
    ])
    para(doc, "Fragmento representativo (job 03) — uso de window functions en SparkSQL:")
    code_block(doc,
        "ROUND(SUM(revenue) OVER (\n"
        "    PARTITION BY category_name\n"
        "    ORDER BY year, month\n"
        "    ROWS BETWEEN UNBOUNDED PRECEDING AND CURRENT ROW\n"
        "), 2) AS cumulative_revenue,\n"
        "RANK() OVER (PARTITION BY year, month ORDER BY revenue DESC) AS revenue_rank_in_month")

    h2(doc, "5.3 Idempotencia")
    para(doc, "Para que el pipeline pueda re-ejecutarse sin errores, el helper drop_table() en "
              "spark/jobs/common.py elimina la tabla Hive y su directorio fisico antes de cada "
              "saveAsTable, evitando el error LOCATION_ALREADY_EXISTS sobre tablas gestionadas.")
    code_block(doc,
        "def drop_table(spark, table):\n"
        "    spark.sql(f'DROP TABLE IF EXISTS {HIVE_DB}.{table}')\n"
        "    tbl_dir = f'/opt/hive/data/warehouse/{HIVE_DB}.db/{table}'\n"
        "    if os.path.exists(tbl_dir):\n"
        "        shutil.rmtree(tbl_dir)")
    para(doc, "Ejecucion manual del pipeline completo:")
    code_block(doc, "make spark-pipeline   # corre los jobs 01 -> 06 en orden")
    screenshot_placeholder(doc, "Spark — evidencia de ejecucion",
        "captura del Spark Master UI (http://localhost:8080) con los jobs completados, o la salida de consola de un job mostrando los conteos de filas.",
        "Docs/screenshots/spark_jobs.png")

    # ── 6. SUPERSET ──────────────────────────────────────────────────────────────
    h1(doc, "6. Visualizacion con Apache Superset")
    para(doc, "Corresponde al componente 3 del enunciado (15%). Superset consume la base "
              "postgres-serving, donde Spark deposito los marts. La configuracion completa "
              "(conexion, datasets, 9 charts y 3 dashboards) esta automatizada en "
              "superset/setup_dashboards.py.")

    h2(doc, "6.1 Automatizacion")
    para(doc, "El script combina la API REST de Superset con su ORM (via create_app) para sortear "
              "limitaciones de la version 4.1.1:")
    bullets(doc, [
        "Autenticacion con requests.Session() para preservar la cookie del token CSRF.",
        "La conexion a la base de datos se crea por ORM (la REST API devolvia 422 al testear la conexion).",
        "Los datasets se crean por REST API y se refrescan sus columnas.",
        "La vinculacion de charts a dashboards se hace por ORM (Dashboard.slices.append).",
        "Los graficos de barras usan el viz_type dist_bar, el unico tipo de barras categorico estable en esta version.",
    ])
    para(doc, "Comandos de puesta en marcha:")
    code_block(doc,
        "make superset-init         # crea admin, db upgrade, superset init\n"
        "make superset-dashboards   # crea conexion, datasets, charts y dashboards")

    h2(doc, "6.2 Dashboards requeridos")
    para(doc, "Los tres dashboards solicitados por el enunciado, con sus graficos:")
    add_table(doc,
        ["Dashboard", "Dataset (mart)", "Graficos"],
        [
            ["1. Ingresos por Mes y Categoria", "mart_revenue_by_cat",
             "Lineas (ingresos por mes/categoria), barras (revenue por categoria), barras temporales (pedidos/unidades)."],
            ["2. Actividad de Clientes por Zona", "mart_orders_by_zone",
             "Barras de pedidos por zona, revenue por zona, distancia y tiempo promedio por zona."],
            ["3. Pedidos Completados vs Cancelados", "mart_order_status",
             "Lineas de pedidos por estado, pie de distribucion de estados, barras de revenue por estado."],
        ],
        col_widths=[2.1, 1.8, 2.9], font_size=8.5, first_col_bold=True)

    h3(doc, "Dashboard 1 — Ingresos por Mes y Categoria")
    screenshot_placeholder(doc, "Dashboard 1",
        "captura completa del dashboard '1. Ingresos por Mes y Categoria' mostrando el grafico de lineas y el de barras por categoria.",
        "Docs/screenshots/dashboard1_ingresos.png")
    h3(doc, "Dashboard 2 — Actividad de Clientes por Zona")
    screenshot_placeholder(doc, "Dashboard 2",
        "captura del dashboard '2. Actividad de Clientes por Zona' con las barras de pedidos y revenue por zona.",
        "Docs/screenshots/dashboard2_zonas.png")
    h3(doc, "Dashboard 3 — Pedidos Completados vs Cancelados")
    screenshot_placeholder(doc, "Dashboard 3",
        "captura del dashboard '3. Pedidos Completados vs Cancelados' con el grafico de lineas y el pie (completed 68%, cancelled 15%).",
        "Docs/screenshots/dashboard3_estados.png")

    # ── 7. AIRFLOW ───────────────────────────────────────────────────────────────
    h1(doc, "7. Orquestacion con Apache Airflow")
    para(doc, "Corresponde al componente 4 del enunciado (15%). El DAG etl_olap_dag programa el "
              "pipeline de forma diaria (@daily) y es modular: cada paso es una tarea reusable.")
    add_image(doc, os.path.join(DIAGRAMS, "airflow_dag.png"), 6.6,
              "Figura 3. Grafo de dependencias del DAG (10 tareas).")

    h2(doc, "7.1 Tareas del DAG")
    add_table(doc,
        ["Tarea", "Tipo", "Funcion"],
        [
            ["extract_postgres", "Python", "Ejecuta el job Spark 01 (extraccion)"],
            ["build_hive_dw", "Python", "Ejecuta el job Spark 02 (construye el DW)"],
            ["run_trends", "Python", "Job Spark 03 (tendencias)"],
            ["run_peak_hours", "Python", "Job Spark 04 (horarios pico)"],
            ["run_monthly_growth", "Python", "Job Spark 05 (crecimiento)"],
            ["check_catalog_changed", "Branch", "Hashea el catalogo y decide si reindexar"],
            ["reindex_elasticsearch", "Python", "Reindexa productos en ES (rama si cambio)"],
            ["skip_reindex", "Empty", "Rama si el catalogo no cambio"],
            ["export_marts", "Python", "Job Spark 06 (exporta marts a serving)"],
            ["data_quality_checks", "Python", "Valida los marts (trigger_rule=all_done)"],
        ],
        col_widths=[2.1, 1.0, 3.5], font_size=8.5, first_col_bold=True)

    h2(doc, "7.2 Patrones de diseno aplicados")
    bullets(doc, [
        ("Ejecucion de Spark via Docker SDK: ", "las tareas hacen exec_run sobre el contenedor "
         "py02_spark_master montando el socket de Docker, evitando reinstalar Spark en Airflow."),
        ("Branching condicional: ", "check_catalog_changed compara el hash MD5 del catalogo entre "
         "ejecuciones (via XCom) y solo dispara el reindexado de Elasticsearch si cambio, "
         "cumpliendo el requisito de 'reindexar si cambia el catalogo'."),
        ("Validacion de calidad: ", "data_quality_checks verifica que los marts tengan filas, que "
         "haya >1000 pedidos y que no existan nulos por encima del limite; usa trigger_rule "
         "'all_done' para correr aunque la rama haya saltado el reindex."),
        ("Paralelismo: ", "los tres analisis y el branch de catalogo se ejecutan en paralelo tras "
         "construir el DW."),
    ])
    para(doc, "Validacion ejecutada (a nivel de tarea):")
    code_block(doc, "airflow tasks test etl_olap_dag <task_id> 2024-01-01")
    screenshot_placeholder(doc, "Airflow UI — DAG ejecutado",
        "captura de la vista Graph o Grid del DAG etl_olap_dag en http://localhost:8081 con una ejecucion exitosa (tareas en verde).",
        "Docs/screenshots/airflow_dag_run.png")

    # ── 8. NEO4J ─────────────────────────────────────────────────────────────────
    h1(doc, "8. Analisis de grafos con Neo4J")
    para(doc, "Corresponde al componente 5 del enunciado (15%). Los datos del OLTP se exportan a "
              "CSV (neo4j/export_csvs.py) y se cargan con LOAD CSV (neo4j/01_load_graph.cypher).")
    add_image(doc, os.path.join(DIAGRAMS, "graph_model.png"), 6.6,
              "Figura 4. Modelo de grafo: 6 tipos de nodo y 7 relaciones.")
    screenshot_placeholder(doc, "Neo4J Browser — grafo cargado",
        "captura del Neo4J Browser (http://localhost:7474) visualizando una muestra del grafo (nodos y relaciones reales).",
        "Docs/screenshots/neo4j_grafo.png")

    h2(doc, "8.1 Nodos y relaciones")
    add_table(doc,
        ["Relacion", "Patron", "Propiedades"],
        [
            ["PLACED", "(User)-[:PLACED]->(Order)", "-"],
            ["FROM_RESTAURANT", "(Order)->(Restaurant)", "-"],
            ["CONTAINS", "(Order)->(Product)", "quantity, unit_price, line_total"],
            ["DELIVERED_TO", "(Order)->(Location)", "distance_km, eta_min, status"],
            ["ASSIGNED", "(Courier)->(Order)", "-"],
            ["RECOMMENDS", "(User)->(User)", "co-clientes del mismo restaurante"],
            ["ROUTE", "(Location)<->(Location)", "distance_km, time_min (bidireccional)"],
        ],
        col_widths=[1.6, 2.5, 2.4], font_size=8.5, first_col_bold=True)

    h2(doc, "8.2 Consultas Cypher requeridas")
    para(doc, "El archivo neo4j/02_queries.cypher implementa las tres consultas del enunciado "
              "mas algoritmos GDS:")
    bullets(doc, [
        ("Top-5 productos comprados juntos: ", "co-ocurrencia de productos en el mismo pedido "
         "(patron (p1)<-[:CONTAINS]-(o)-[:CONTAINS]->(p2))."),
        ("Usuarios que recomiendan a otros: ", "por grado de salida de RECOMMENDS y, ademas, "
         "PageRank con GDS para medir influencia."),
        ("Caminos minimos entre ubicaciones: ", "shortestPath nativo y Dijkstra ponderado por "
         "distancia con GDS, para reparto eficiente."),
    ])
    para(doc, "Ejemplo — top-5 productos comprados juntos:")
    code_block(doc,
        "MATCH (p1:Product)<-[:CONTAINS]-(o:Order)-[:CONTAINS]->(p2:Product)\n"
        "WHERE p1.id < p2.id\n"
        "WITH p1.name AS producto_1, p2.name AS producto_2,\n"
        "     COUNT(DISTINCT o) AS pedidos_juntos\n"
        "ORDER BY pedidos_juntos DESC LIMIT 5\n"
        "RETURN producto_1, producto_2, pedidos_juntos;")
    para(doc, "Comandos:")
    code_block(doc,
        "make neo4j-export-csv   # exporta CSVs desde el OLTP\n"
        "make neo4j-load         # carga el grafo\n"
        "make neo4j-queries      # ejecuta las consultas Cypher")
    screenshot_placeholder(doc, "Neo4J — resultado de consulta",
        "captura del resultado de la consulta de top-5 productos comprados juntos (o usuarios influyentes por PageRank) en el Neo4J Browser.",
        "Docs/screenshots/neo4j_consulta.png")

    # ── 9. ROUTING ───────────────────────────────────────────────────────────────
    h1(doc, "9. Asignacion de rutas de entrega")
    para(doc, "Corresponde al componente 6 del enunciado (10%). El modulo "
              "routing/route_assignment.py simula la asignacion de pedidos pendientes a "
              "repartidores y optimiza el orden de visita.")
    bullets(doc, [
        ("Heuristica: ", "vecino mas cercano (nearest-neighbor) sobre los pedidos pendientes, "
         "repartiendo la carga equitativamente entre repartidores."),
        ("Distancias: ", "se obtienen del grafo de zonas con Dijkstra (GDS); si GDS falla, hace "
         "fallback a la distancia Haversine entre coordenadas."),
        ("Salida: ", "imprime las rutas por repartidor y exporta routing/route_results.json con "
         "pedidos asignados, distancia total y secuencia de entrega."),
    ])
    para(doc, "Ejecucion:")
    code_block(doc, "make neo4j-routing   # corre el algoritmo y genera route_results.json")
    screenshot_placeholder(doc, "Enrutamiento — salida",
        "captura de la salida en consola de route_assignment.py (rutas por repartidor) o un extracto de routing/route_results.json.",
        "Docs/screenshots/routing_salida.png")

    # ── 10. DECISIONES DE DISENO ─────────────────────────────────────────────────
    h1(doc, "10. Decisiones de diseno")
    add_table(doc,
        ["Decision", "Alternativas", "Justificacion"],
        [
            ["Hive como DW (no solo Postgres)", "Postgres analitico, DuckDB",
             "Cumple el requisito explicito de DW open-source con Hive; modelo dimensional escalable."],
            ["Metastore Hive sobre Postgres", "Derby embebido",
             "Derby es monousuario y fragil; Postgres permite concurrencia y persistencia."],
            ["ORC + Snappy", "Parquet, texto",
             "Columnar, buena compresion, predicate pushdown nativo en Hive/Spark."],
            ["Particionar hechos por year/month", "Sin particion",
             "Partition pruning acelera las consultas temporales, que son la mayoria."],
            ["Capa serving en Postgres aparte", "Superset directo a Hive",
             "Superset consulta SQL estandar rapido sin depender del cliente Hive/Thrift."],
            ["DataFrame API + SparkSQL mezclados", "Solo uno",
             "Requisito del enunciado; SQL para ventanas/joins legibles, API para transformaciones."],
            ["Airflow ejecuta Spark via Docker SDK", "SparkSubmitOperator",
             "Evita instalar Spark dentro de Airflow; reusa el contenedor spark-master existente."],
            ["Perfiles de Docker Compose", "Un unico stack",
             "Permite levantar solo lo necesario por fase y ahorrar memoria."],
            ["Superset configurado por ORM", "Solo REST API",
             "La REST API 4.1.1 falla al crear la conexion (422) y al vincular charts; el ORM es estable."],
            ["Datos sinteticos reproducibles", "Datos reales del P1",
             "El P1 no tiene historial transaccional; semilla fija garantiza reproducibilidad."],
        ],
        col_widths=[1.9, 1.6, 3.0], font_size=8.5, first_col_bold=True)

    # ── 11. PRUEBAS ──────────────────────────────────────────────────────────────
    h1(doc, "11. Pruebas y validaciones")
    bullets(doc, [
        ("Integridad del warehouse: ", "los jobs Spark imprimen el conteo de filas de cada "
         "dimension y hecho; data_quality_checks valida en serving que existan filas y >1000 "
         "pedidos sin nulos por encima del limite."),
        ("Pipeline Airflow: ", "validado a nivel de tarea con 'airflow tasks test' para las 10 "
         "tareas, confirmando ejecucion de extremo a extremo (no solo el parseo del DAG)."),
        ("Spark y Neo4J: ", "los jobs muestran resultados (show) y las consultas Cypher devuelven "
         "rankings; el routing exporta un JSON inspeccionable."),
        ("Visualizacion: ", "los 3 dashboards renderizan correctamente sus 9 graficos tras migrar "
         "las barras a dist_bar."),
    ])

    # ── 12. DESPLIEGUE ───────────────────────────────────────────────────────────
    h1(doc, "12. Instrucciones de despliegue y uso")

    h2(doc, "12.1 Prerrequisitos")
    bullets(doc, [
        "Docker Desktop con Docker Compose v2.",
        "Python 3.11+ en el host (para el generador de datos y los scripts de Superset/routing).",
        "Recomendado: 8 GB de RAM libres para levantar el stack completo.",
    ])

    h2(doc, "12.2 Levantamiento por fases (recomendado)")
    code_block(doc,
        "# 1. Fuentes y generacion de datos\n"
        "make core           # postgres-oltp, mongo, elasticsearch\n"
        "make seed           # genera 20k pedidos en el OLTP\n\n"
        "# 2. Data Warehouse + Spark\n"
        "make dw             # hive metastore, hiveserver2, spark master/worker\n"
        "make hive-init      # crea DB, dimensiones, hechos y cubos\n"
        "make spark-pipeline # ETL completo (jobs 01-06)\n\n"
        "# 3. Orquestacion\n"
        "make airflow-up     # Airflow UI en http://localhost:8081\n\n"
        "# 4. Grafos y rutas\n"
        "make neo4j-export-csv && make neo4j-load\n"
        "make neo4j-queries && make neo4j-routing\n\n"
        "# 5. Visualizacion\n"
        "make viz                  # postgres-serving + superset\n"
        "make superset-init        # admin / init\n"
        "make superset-dashboards  # crea los 3 dashboards")

    h2(doc, "12.3 Levantamiento completo y apagado")
    code_block(doc,
        "make up-all   # levanta TODOS los perfiles a la vez\n"
        "make ps       # estado de los contenedores\n"
        "make down     # detiene el stack\n"
        "make clean    # detiene y elimina volumenes (borra datos)")

    # ── 13. COBERTURA DE REQUISITOS ──────────────────────────────────────────────
    h1(doc, "13. Cobertura de requisitos (rubrica)")
    add_table(doc,
        ["Criterio de evaluacion", "Peso", "Evidencia", "Estado"],
        [
            ["Data Warehouse y OLAP", "20%", "warehouse/*.hql (8 dim, 2 hechos, 6 cubos)", "Completo"],
            ["Procesamiento con Spark", "20%", "spark/jobs/*.py (6 jobs, 3 analisis)", "Completo"],
            ["Visualizacion", "15%", "superset/setup_dashboards.py (3 dashboards)", "Completo"],
            ["Airflow", "15%", "airflow/dags/etl_olap_dag.py (10 tareas)", "Completo"],
            ["Neo4J", "15%", "neo4j/*.cypher (consultas + GDS)", "Completo"],
            ["Enrutamiento", "10%", "routing/route_assignment.py", "Completo"],
            ["Documentacion", "5%", "Este documento + diagramas", "Completo"],
        ],
        col_widths=[2.3, 0.6, 2.7, 0.9], font_size=8.5, first_col_bold=True)
    para(doc, "Entregables del enunciado:")
    add_table(doc,
        ["Entregable", "Ubicacion"],
        [
            ["Codigo fuente y DAG de Airflow", "spark/, warehouse/, airflow/dags/"],
            ["Scripts/notebooks de Spark", "spark/jobs/01..06"],
            ["Dashboards exportables", "superset/setup_dashboards.py (reproducible)"],
            ["Consultas Cypher y estructura del grafo", "neo4j/01_load_graph.cypher, 02_queries.cypher"],
            ["Capturas / video demostrativo", "Docs/screenshots/ (placeholders en secciones 4, 5, 6, 7, 8, 9)"],
            ["Documentacion tecnica en PDF", "Exportar este .docx a PDF"],
        ],
        col_widths=[3.0, 3.6], font_size=9, first_col_bold=True)

    # ── 14. ESTRUCTURA ───────────────────────────────────────────────────────────
    h1(doc, "14. Estructura del repositorio")
    code_block(doc,
        "Proyecto 2 OLAP/\n"
        "  docker-compose.yml        # 5 perfiles, 11 servicios\n"
        "  Makefile                  # targets por fase\n"
        "  .env                      # config y credenciales\n"
        "  data-generator/           # generacion de datos sinteticos (OLTP)\n"
        "  warehouse/                # DDL Hive: dimensiones, hechos, cubos\n"
        "  spark/jobs/               # 6 jobs ETL + common.py\n"
        "  airflow/                  # Dockerfile + dags/etl_olap_dag.py\n"
        "  neo4j/                    # export_csvs, load_graph, queries, import/\n"
        "  routing/                  # route_assignment.py + route_results.json\n"
        "  superset/                 # Dockerfile + setup_dashboards.py\n"
        "  Docs/                     # diagramas, capturas y esta documentacion\n"
        "  Proyecto1/                # referencia de solo lectura (no se modifica)",
        size=8.5)

    para(doc, "")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— Fin del documento —")
    r.italic = True; r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUT_DOCX)
    print(f"\nDocumento generado: {OUT_DOCX}")


# ═════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Renderizando diagramas...")
    diagram_architecture()
    diagram_star_schema()
    diagram_graph_model()
    diagram_airflow_dag()
    print("Construyendo documento Word...")
    build()
